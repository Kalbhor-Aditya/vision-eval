"""Generate VisionEval end-to-end documentation as a .docx file."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys
from pathlib import Path

OUT_PATH = Path(__file__).parent.parent / "VisionEval_Documentation.docx"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_code(doc, code):
    p = doc.add_paragraph(style="No Spacing")
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    p.paragraph_format.left_indent = Inches(0.3)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()
    return table


def build_doc():
    doc = Document()

    # ── Title ────────────────────────────────────────────────────────────────
    title = doc.add_heading("VisionEval: Multimodal Reasoning & Evaluation Platform", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("End-to-End Technical Documentation").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── 1. Project Overview ───────────────────────────────────────────────────
    add_heading(doc, "1. Project Overview")
    add_para(doc, (
        "VisionEval is a production-grade multimodal AI platform that combines open-source vision "
        "large language models (VLMs) with a structured skills architecture, comprehensive evaluation "
        "harness, and deep observability via Langfuse. The system goes beyond basic image captioning "
        "by decomposing visual reasoning tasks into specialized skills, tracing every inference, and "
        "providing automated + human evaluation of every response."
    ))

    add_heading(doc, "Key Differentiators", level=2)
    points = [
        "Skills-based architecture: agent decomposes any question into 1-2 specialized visual skills",
        "LLM router: a fast text model selects the right skill(s) automatically",
        "Langfuse v4 integration: full tracing, scoring, dataset management, and prompt versioning",
        "Three automated eval types: LLM-as-judge, semantic similarity, consistency checking",
        "CI-style eval harness with configurable pass/fail thresholds per skill",
        "Structured logging: rich terminal output + rotating JSON log files",
        "No local model storage: all inference via Groq API (open-source models)",
    ]
    for p in points:
        doc.add_paragraph(p, style="List Bullet")

    # ── 2. Architecture ───────────────────────────────────────────────────────
    add_heading(doc, "2. System Architecture")
    add_para(doc, (
        "The system follows a layered architecture with clear separation between the API layer, "
        "skills engine, evaluation pipeline, and observability stack."
    ))

    add_heading(doc, "2.1 High-Level Flow", level=2)
    add_para(doc, "A user request flows through the system as follows:")
    flow = [
        "1. User uploads image + question via Streamlit UI",
        "2. Streamlit sends POST /analyze to FastAPI backend",
        "3. @observe decorator creates a Langfuse trace automatically",
        "4. Skill Router (Groq llama-3.1-8b-instant) selects 1-2 skills",
        "5. Each skill calls vision_query() → Groq Llama-4 Scout",
        "6. Skill returns structured SkillResult (observations, reasoning, answer)",
        "7. If multiple skills: synthesizer merges answers",
        "8. LLM Judge scores the final answer (0–10)",
        "9. Score written to Langfuse trace automatically",
        "10. Response returned to UI; user can add human feedback",
        "11. All events logged to terminal (rich) + logs/app.jsonl",
    ]
    for step in flow:
        doc.add_paragraph(step, style="List Number")

    add_heading(doc, "2.2 Directory Structure", level=2)
    add_code(doc, """Vision/
├── app/
│   ├── main.py              # FastAPI endpoints
│   ├── config.py            # Pydantic settings (env vars)
│   ├── logger.py            # Centralized logging setup
│   ├── models.py            # Pydantic request/response schemas
│   ├── vlm_client.py        # Groq + HuggingFace API wrappers
│   ├── skill_registry.py    # Skill registry + LLM router
│   ├── langfuse_client.py   # Langfuse v4 helpers (@observe wrappers)
│   ├── skills/
│   │   ├── base.py          # BaseSkill (abstract, @observe, CoT parsing)
│   │   ├── read_text.py     # OCR skill
│   │   ├── detect_objects.py
│   │   ├── analyze_chart.py
│   │   ├── describe_scene.py
│   │   ├── verify_claim.py
│   │   ├── compare_images.py  # Dual-image comparison
│   │   └── extract_structured.py
│   └── evals/
│       ├── llm_judge.py     # LLM-as-judge (Groq)
│       ├── semantic_sim.py  # Sentence-transformer cosine similarity
│       └── consistency.py  # Multi-rephrase consistency check
├── harness/
│   ├── runner.py            # CLI harness runner
│   ├── suite_definitions.py # Test cases per skill
│   ├── thresholds.py        # Pass/fail thresholds
│   └── report.py            # Rich terminal + JSONL reports
├── ui/
│   └── streamlit_app.py     # 4-page Streamlit UI
├── scripts/
│   ├── seed_dataset.py      # Seed Langfuse datasets
│   └── generate_docs.py     # This documentation generator
└── logs/                    # Auto-created log files""")

    # ── 3. Technology Stack ───────────────────────────────────────────────────
    add_heading(doc, "3. Technology Stack")
    add_table(doc,
        ["Layer", "Technology", "Purpose"],
        [
            ["Vision Model", "Groq: meta-llama/llama-4-scout-17b-16e-instruct", "Primary vision inference (free API)"],
            ["Text Model", "Groq: llama-3.1-8b-instant", "Skill routing, LLM judge, synthesis"],
            ["Vision Fallback", "HuggingFace: Qwen/Qwen2-VL-7B-Instruct", "Auto-fallback if Groq vision fails"],
            ["Observability", "Langfuse v4", "Tracing, scoring, datasets, prompt mgmt"],
            ["Semantic Eval", "sentence-transformers (all-MiniLM-L6-v2)", "Cosine similarity scoring"],
            ["API Backend", "FastAPI + uvicorn", "Async REST API"],
            ["UI", "Streamlit", "4-page web interface"],
            ["Config", "pydantic-settings", "Environment variable management"],
            ["Logging", "Python logging + rich + RotatingFileHandler", "Terminal + file logging"],
        ]
    )

    # ── 4. Skills System ──────────────────────────────────────────────────────
    add_heading(doc, "4. Skills Architecture")
    add_para(doc, (
        "The skills system is the core innovation of VisionEval. Instead of asking the VLM one generic "
        "question, the system routes the query to specialized skills — each with a focused system prompt "
        "optimized for a specific type of visual reasoning task."
    ))

    add_heading(doc, "4.1 Available Skills", level=2)
    add_table(doc,
        ["Skill", "Use Case", "Example Query"],
        [
            ["read_text", "OCR — extract all visible text", "What does the sign say?"],
            ["detect_objects", "Identify + count objects/entities", "How many people are in this image?"],
            ["analyze_chart", "Interpret charts, graphs, tables", "What is the trend in this graph?"],
            ["describe_scene", "Rich scene description", "Describe what is happening here"],
            ["verify_claim", "Fact-check a claim against the image", "Is the claim X supported by this image?"],
            ["compare_images", "Diff two images (dual-image input)", "What changed between these two screenshots?"],
            ["extract_structured", "Extract key-value structured data", "Extract all fields from this invoice"],
        ]
    )

    add_heading(doc, "4.2 Chain-of-Thought (CoT) Format", level=2)
    add_para(doc, (
        "Every skill appends a CoT suffix to its prompt, forcing the model to structure its response "
        "in three sections. This makes reasoning inspectable and parseable:"
    ))
    add_code(doc, """OBSERVATIONS: <what you see in the image>
REASONING: <step-by-step thinking>
ANSWER: <final answer>""")

    add_heading(doc, "4.3 Skill Router", level=2)
    add_para(doc, (
        "The router uses Groq's fast text model to select 1-2 appropriate skills for a given question. "
        "It receives the skill registry (names + descriptions) as context and responds with a JSON array "
        "of skill names. If parsing fails, it defaults to 'describe_scene'."
    ))
    add_code(doc, """# Router system prompt includes all skill descriptions
# Model responds with: ["read_text", "verify_claim"]
# Invalid skill names are filtered; fallback to ["describe_scene"]""")

    add_heading(doc, "4.4 BaseSkill and @observe", level=2)
    add_para(doc, (
        "All skills inherit from BaseSkill. The invoke() method is decorated with Langfuse's @observe "
        "decorator, which automatically creates a child span in the current trace. The skill updates "
        "the observation with input/output metadata, enabling full tracing in Langfuse."
    ))
    add_code(doc, """class BaseSkill(ABC):
    @observe()
    async def invoke(self, image_b64: str, question: str) -> SkillResult:
        update_observation(name=self.name, input={"question": question})
        raw, meta = await vision_query(image_b64, prompt, system_prompt)
        obs, reasoning, answer = parse_cot(raw)
        update_observation(output={"answer": answer}, metadata={...})
        return SkillResult(...)""")

    # ── 5. Langfuse Integration ───────────────────────────────────────────────
    add_heading(doc, "5. Langfuse Integration")
    add_para(doc, (
        "Langfuse is the observability backbone of VisionEval. Every aspect of the system — from "
        "individual VLM calls to eval scores to human feedback — is captured in Langfuse traces."
    ))

    add_heading(doc, "5.1 Langfuse v4 Decorator Pattern", level=2)
    add_para(doc, (
        "Langfuse v4 replaced the manual trace/span construction API with the @observe decorator. "
        "Any function decorated with @observe automatically becomes a span in the current trace. "
        "Nested @observe calls create a parent-child span hierarchy."
    ))
    add_code(doc, """from langfuse.decorators import observe, langfuse_context

@observe(name="analyze")          # Creates root trace
async def _run_analyze(req):
    trace_id = langfuse_context.get_current_trace_id()

    # Skill.invoke() is also @observe → becomes child span
    result = await skill.invoke(image_b64, question)

    # Score the trace from within the context
    langfuse_context.score_current_trace(name="judge_score", value=0.8)""")

    add_heading(doc, "5.2 What Gets Traced", level=2)
    add_table(doc,
        ["Langfuse Feature", "How Used", "Data Captured"],
        [
            ["Traces", "@observe on _run_analyze, _run_compare", "Full request lifecycle"],
            ["Spans", "@observe on each skill.invoke()", "Per-skill latency, input, output"],
            ["Observations", "update_observation() calls", "Model, tokens, latency, answer"],
            ["Scores (auto)", "score_current_trace() after eval", "judge_score normalized 0–1"],
            ["Scores (human)", "add_score() in /feedback endpoint", "User 0–1 rating + comment"],
            ["Datasets", "save_to_dataset() in harness", "Test inputs + expected outputs"],
            ["flush()", "After each request", "Ensures data is sent to Langfuse"],
        ]
    )

    add_heading(doc, "5.3 Scoring from Outside @observe Context", level=2)
    add_para(doc, (
        "The /feedback endpoint receives a trace_id and score after the fact (outside any @observe "
        "context). It uses lf.score(trace_id=...) directly to attach the human score to the original trace."
    ))

    # ── 6. Evaluation System ──────────────────────────────────────────────────
    add_heading(doc, "6. Evaluation System")

    add_heading(doc, "6.1 LLM-as-Judge (app/evals/llm_judge.py)", level=2)
    add_para(doc, (
        "A second LLM (Groq llama-3.1-8b-instant) is used as an impartial judge to score every "
        "VLM response. The judge receives the original question and model answer, then scores it "
        "0–10 on accuracy, helpfulness, and reasoning quality. The score is normalized to 0–1 "
        "and written to Langfuse as 'judge_score'."
    ))
    add_code(doc, """# Judge prompt asks for: SCORE: <0-10>\nREASON: <one sentence>
score, reason = await judge(question="...", answer="...")
# Returns: (7.5, "Answer correctly identifies the main objects")""")

    add_heading(doc, "6.2 Semantic Similarity (app/evals/semantic_sim.py)", level=2)
    add_para(doc, (
        "Uses the all-MiniLM-L6-v2 sentence-transformer model to compute cosine similarity "
        "between a model answer and a reference answer. Returns a float in [0, 1]. "
        "Used when ground-truth answers are available (e.g., in harness test cases)."
    ))

    add_heading(doc, "6.3 Consistency Check (app/evals/consistency.py)", level=2)
    add_para(doc, (
        "The consistency eval detects whether a model gives contradictory answers to semantically "
        "equivalent questions. It works by: (1) generating 2 rephrased versions of the question, "
        "(2) running the same skill on all 3 questions, (3) computing pairwise semantic similarity "
        "of the 3 answers. Low consistency scores indicate model instability."
    ))

    # ── 7. Eval Harness ───────────────────────────────────────────────────────
    add_heading(doc, "7. Eval Harness")
    add_para(doc, (
        "The harness is a standalone CLI runner that executes test suites against all skills. "
        "It is designed like a CI test runner — each test case has a pass/fail verdict based "
        "on configurable score thresholds."
    ))

    add_heading(doc, "7.1 Running the Harness", level=2)
    add_code(doc, """# Run all suites
python -m harness.runner --suite all

# Run a specific skill suite
python -m harness.runner --suite read_text

# Push results to Langfuse datasets
python -m harness.runner --suite all --push""")

    add_heading(doc, "7.2 Thresholds", level=2)
    add_table(doc,
        ["Skill", "judge_score threshold (0–10)"],
        [
            ["read_text", "7.0"],
            ["detect_objects", "6.5"],
            ["analyze_chart", "6.0"],
            ["describe_scene", "6.0"],
            ["verify_claim", "7.0"],
            ["extract_structured", "7.0"],
            ["compare_images", "6.0"],
        ]
    )

    add_heading(doc, "7.3 Output", level=2)
    add_para(doc, (
        "The harness produces: (1) a rich colored table in the terminal showing each test case's "
        "score and pass/fail status, (2) a JSON line appended to logs/harness.jsonl with the full "
        "run summary. Optionally pushes each result to Langfuse as a dataset item."
    ))

    # ── 8. Logging ────────────────────────────────────────────────────────────
    add_heading(doc, "8. Logging Architecture")
    add_para(doc, (
        "All modules import from a single centralized logger (app/logger.py). "
        "Three handlers are active simultaneously:"
    ))
    add_table(doc,
        ["Handler", "Output", "Format", "Level"],
        [
            ["RichHandler", "Terminal (stdout)", "Colored, human-readable", "Configurable (default INFO)"],
            ["RotatingFileHandler", "logs/app.jsonl (10MB, 5 backups)", "JSON lines (structured)", "Configurable"],
            ["RotatingFileHandler", "logs/errors.log (5MB, 3 backups)", "Plain text", "ERROR only"],
        ]
    )

    add_heading(doc, "Log Event Examples", level=2)
    add_code(doc, """# Every log call includes structured extra fields:
logger.info("Skill invoked", extra={"skill": "read_text", "question": "..."})
logger.info("VLM response", extra={"model": "llama-4-scout", "tokens": 312, "latency_ms": 1240})
logger.info("Trace scored", extra={"score_name": "judge_score", "score_value": 0.75})
logger.error("Skill failed", extra={"skill": "analyze_chart", "error": "timeout"})""")

    # ── 9. API Reference ──────────────────────────────────────────────────────
    add_heading(doc, "9. API Reference")
    add_table(doc,
        ["Endpoint", "Method", "Purpose"],
        [
            ["GET /health", "GET", "Health check"],
            ["GET /skills", "GET", "List all available skills with descriptions"],
            ["POST /analyze", "POST", "Analyze an image with optional skill selection"],
            ["POST /compare", "POST", "Compare two images"],
            ["POST /feedback", "POST", "Submit human feedback score for a trace"],
        ]
    )

    add_heading(doc, "/analyze Request Body", level=2)
    add_code(doc, """{
  "image_b64": "<base64-encoded JPEG>",
  "question": "What does this receipt say?",
  "skill": "auto",          // or: "read_text", "detect_objects", etc.
  "session_id": "optional-session-id"
}""")

    add_heading(doc, "/analyze Response", level=2)
    add_code(doc, """{
  "trace_id": "abc123",
  "skills_used": ["read_text", "verify_claim"],
  "skill_results": [
    {
      "skill_name": "read_text",
      "observations": "...",
      "reasoning": "...",
      "answer": "Total: $42.50",
      "confidence": 0.5,
      "latency_ms": 1200,
      "tokens_used": 280
    }
  ],
  "final_answer": "The receipt shows a total of $42.50.",
  "scores": {"judge_score": 0.85},
  "total_latency_ms": 2100
}""")

    # ── 10. UI Pages ──────────────────────────────────────────────────────────
    add_heading(doc, "10. Streamlit UI Pages")
    add_table(doc,
        ["Page", "Features"],
        [
            ["Analyze", "Upload image, select skill, see CoT chain, view scores, submit human feedback"],
            ["Compare", "Upload two images, ask comparison question, see side-by-side descriptions"],
            ["Harness", "Run eval harness from UI, view pass/fail table, see recent harness run history"],
            ["Logs", "Live log viewer with level filtering (tail logs/app.jsonl)"],
        ]
    )

    # ── 11. Setup & Running ───────────────────────────────────────────────────
    add_heading(doc, "11. Setup & Running")

    add_heading(doc, "11.1 Prerequisites", level=2)
    for item in [
        "Python 3.11+",
        "Groq API key: free at https://console.groq.com",
        "Langfuse account: free at https://cloud.langfuse.com (get public + secret keys)",
        "HuggingFace token (optional fallback): https://huggingface.co/settings/tokens",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "11.2 Installation & Start", level=2)
    add_code(doc, """cp .env.example .env
# Edit .env: fill GROQ_API_KEY, LANGFUSE_PUBLIC_KEY, LANGFUSE_SECRET_KEY

pip install -r requirements.txt

# Terminal 1: Start API
uvicorn app.main:app --reload

# Terminal 2: Start UI
streamlit run ui/streamlit_app.py

# Terminal 3: Run eval harness
python -m harness.runner --suite all""")

    add_heading(doc, "11.3 Verify Everything Works", level=2)
    steps = [
        "Open http://localhost:8501 → Streamlit UI loads",
        "Go to Analyze page → upload an image → ask a question → see skill chain + answer",
        "Open Langfuse dashboard → find the trace → verify skill spans and judge_score",
        "Submit human feedback → verify score appears in Langfuse",
        "Run harness → see pass/fail table → check logs/harness.jsonl",
        "Go to Logs page → see structured log entries",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}", style="List Number")

    # ── 12. Key Concepts Glossary ─────────────────────────────────────────────
    add_heading(doc, "12. Key Concepts Glossary")
    add_table(doc,
        ["Concept", "Definition"],
        [
            ["Skill", "A focused VLM capability with a specialized system prompt and typed output"],
            ["Skill Router", "LLM that selects which skill(s) to invoke for a given query"],
            ["Chain-of-Thought", "Structured prompt format: OBSERVATIONS → REASONING → ANSWER"],
            ["@observe", "Langfuse v4 decorator that auto-creates a trace/span for a function"],
            ["Langfuse Trace", "Root-level record of one user request + all its child spans"],
            ["Langfuse Span", "Child record within a trace (one per skill invocation)"],
            ["LLM-as-Judge", "Using a second LLM to score the quality of the first LLM's answer"],
            ["Consistency Eval", "Testing if a model gives stable answers to rephrased questions"],
            ["Eval Harness", "Automated test runner with predefined cases + pass/fail thresholds"],
            ["Semantic Similarity", "Cosine similarity of sentence embeddings to compare answer quality"],
        ]
    )

    doc.save(OUT_PATH)
    print(f"Documentation saved to: {OUT_PATH}")


if __name__ == "__main__":
    build_doc()
